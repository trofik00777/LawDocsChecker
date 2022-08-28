import os
import io

from loguru import logger
from fastapi import Depends, FastAPI, HTTPException, status, File, UploadFile
from fastapi.staticfiles import StaticFiles
from fastapi.security import HTTPBearer
import docx
import subprocess

from .api_models import ResponseModel
from ml.model_manager import DocManager
from ml.models import BaseModel

DocApi = FastAPI()
token_auth_scheme = HTTPBearer()
manager = DocManager()
model = BaseModel()

DocApi.mount("/static", StaticFiles(directory="static"), name="static")


@DocApi.post("/post_file")
async def process_the_document(item: UploadFile = File(...),
                               token: str = Depends(token_auth_scheme)
                               ):
    check_authorization(token)

    try:
        contents = await item.read()
        doc = docx.Document(io.BytesIO(contents))
    except Exception:
        try:
            contents = await item.read()
            with open("./tmp/tmpfile.doc", 'rb') as f:
                f.write(contents)
            subprocess.call(['soffice', '--headless', '--convert-to', 'docx', '--outdir', 'tmp', "./tmp/tmpfile.doc"])
            doc = docx.Document("./tmp/tmpfile.docx")
        except Exception:
            raise HTTPException(
                status_code=status.HTTP_451_UNAVAILABLE_FOR_LEGAL_REASONS,
                detail="File is invalid",
            )
    finally:
        await item.close()

    logger.info('Parsing file')
    result_parsing = manager.parsing_with_brackets(doc)
    # logger.debug(f'{"".join(result_parsing)}')
    info = ResponseModel()
    info.parts = result_parsing

    answer = {"classes": []}
    # print(info.parts)

    for text, flag in info.parts:
        answer["classes"].append({"text": text, "label": model(text) + 1 if flag else -1})

    print(answer)  # TODO: fix

    return answer


@DocApi.post("/post_inline_file")
async def process_the_document_inline(item: UploadFile = File(...),
                               token: str = Depends(token_auth_scheme)
                               ):
    check_authorization(token)
    try:
        contents = await item.read()
    except Exception:
        raise HTTPException(
            status_code=status.HTTP_451_UNAVAILABLE_FOR_LEGAL_REASONS,
            detail="File is invalid",
        )


    try:
        print(contents)
        doc = docx.Document(io.BytesIO(contents))
    except Exception:
        try:
            with open("./tmp_conv/tmpfile.doc", 'wb') as f:
                f.write(contents)
            subprocess.call(['soffice', '--headless', '--convert-to', 'docx', '--outdir', 'tmp_conv', "./tmp_conv/tmpfile.doc"])
            doc = docx.Document("./tmp_conv/tmpfile.docx")

            logger.info("convert doc -> docx")
        except Exception as e:
            logger.error(e)

            raise HTTPException(
                status_code=status.HTTP_451_UNAVAILABLE_FOR_LEGAL_REASONS,
                detail="File is invalid",
            )
    finally:
        await item.close()

    response = manager.colorize_doc_inline(doc)

    return response



def check_authorization(token: str) -> None:
    if token.credentials != os.getenv('AUTHORIZATION_TOKEN'):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication credentials",
            headers={"WWW-Authenticate": "Bearer"},
        )
