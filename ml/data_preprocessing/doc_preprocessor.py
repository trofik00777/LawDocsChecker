import re
from string import punctuation
from typing import List

from loguru import logger

from docx import Document
from docx.shared import RGBColor

import nltk
from nltk.corpus import stopwords

from ml.models import BaseModel

model_base = BaseModel()
TOKENS_BY_NUMBER_LABEL = {}
COLORS = [
    (94, 217, 219),
    (213, 99, 210),
    (113, 123, 130),
    (154, 164, 98),
    (236, 255, 215),
    (61, 165, 59),
    (167, 208, 18),
    (200, 56, 216),
    (239, 4, 10),
    (41, 146, 1),
    (237, 224, 224),
    (239, 33, 244),
    (69, 172, 35),
    (37, 158, 250),
    (117, 254, 156),
    (35, 196, 109),
    (117, 71, 249),
    (142, 184, 7),
    (230, 142, 176),
    (209, 107, 221),
    (150, 184, 138),
    (36, 200, 149),
    (21, 248, 44),
    (134, 8, 45),
    (23, 192, 53),
    (137, 115, 23),
    (144, 70, 194),
    (48, 251, 51),
    (241, 169, 219),
    (181, 255, 191),
    (72, 98, 177),
    (86, 233, 13),
    (174, 31, 176),
    (17, 95, 53),
    (118, 103, 35),
    (117, 99, 65),
    (133, 249, 66),
    (67, 216, 202),
    (205, 241, 67),
    (45, 206, 41)
]

nltk.download("stopwords")

russian_stopwords = stopwords.words("russian")

patterns = re.compile(r"[A-Za-z0-9!#$%&'()*+,./:;<=>?@[\]^_`{|}~—\"\-]+")
split_regex = re.compile(r'[.|!|?|…]')

COUNT_POSSIBLE_CHARS_BETWEEN_BRACKETS = 2


class DocProcessor:
    @staticmethod
    def get_full_text(doc: Document):
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        return '\n'.join(full_text)

    @classmethod
    def preprocess_doc_splitted_by_sentences(cls, doc: Document) -> List[str]:
        sentences = filter(lambda t: t, [t.strip() for t in split_regex.split(cls.get_full_text(doc))])
        sentences = list(map(cls._preprocess_text, sentences))
        sentences = list(filter(lambda elem: not (elem is None), sentences))
        return sentences

    @staticmethod
    def get_paragraphs(text: str) -> List[tuple[str, bool]]:
        is_right = False
        left_bracket = right_bracket = 0
        paragraphs = []
        for i in range(len(text)):
            if text[i] == "{":
                last_i = i
                try:
                    last_i = i + text[i:i + COUNT_POSSIBLE_CHARS_BETWEEN_BRACKETS + 2].index("}")
                except Exception as e:
                    logger.error(f"There are no bracket at '{text[i:i + 10]}'")

                if is_right:
                    right_bracket = i
                    paragraphs.append((text[left_bracket + 1:right_bracket], True))
                else:
                    left_bracket = last_i
                    paragraphs.append((text[right_bracket:left_bracket + 1], False))
                is_right = not is_right
        return paragraphs

    @classmethod
    def preprocess_doc_splitted_by_brackets(cls, doc: Document) -> List[tuple[str, bool]]:
        text = cls.get_full_text(doc)
        logger.debug(text)
        return cls.get_paragraphs(text)

    @staticmethod
    def _preprocess_text(text: str) -> str:
        text = re.sub(patterns, ' ', text)
        tokens = text.lower().split()
        f_tokens = []

        for token in tokens:
            if all([token not in russian_stopwords,
                    token != " ",
                    token.strip() not in punctuation,
                    ]):
                f_tokens.append(token)
        text = " ".join(tokens)
        if len(tokens) > 2:
            return text


class InlineDocProcessor:
    num_file: int = 0

    @classmethod
    def process(cls, doc: Document):
        json = {"classes": []}
        for paragraph in doc.paragraphs:
            if "{" in paragraph.text:
                parts = DocProcessor.get_paragraphs(paragraph.text)
                paragraph.text = ""
                for txt, is_classifier in parts:
                    print(txt)
                    print("-" * 100)
                    if is_classifier:
                        label = model_base(txt) + 1
                        run = paragraph.add_run(txt)
                        run.font.color.rgb = RGBColor(*COLORS[label - 1])

                        json["classes"].append({"text": txt, "label": label})
                    else:
                        paragraph.add_run(txt)

                        json["classes"].append({"text": txt, "label": -1})
            else:
                json["classes"].append({"text": paragraph.text, "label": -1})
        doc.save(f"./static/{cls.num_file}.docx")
        json["filename"] = f"{cls.num_file}.docx"
        cls.num_file += 1

        return json




if __name__ == "__main__":
    file = Document("/Users/trofik00777/Downloads/test.doc")
    # print(DocProcessor.preprocess_doc_splitted_by_brackets(file))
    InlineDocProcessor.process(doc=file)
